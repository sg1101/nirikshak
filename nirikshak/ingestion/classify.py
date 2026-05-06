"""Document routing and ingestion orchestrator."""

import logging
import shutil
from pathlib import Path
from uuid import uuid4

from nirikshak.core.config import get_settings
from nirikshak.core.schemas import Document, Page, RoutingTag
from nirikshak.ingestion.pdf_text import extract_pdf
from nirikshak.ingestion.ocr import ocr_document_pages
from nirikshak.ingestion.vision import vision_page

logger = logging.getLogger(__name__)

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".tif"}


def _store_file(file_path: Path) -> Path:
    """Copy file to persistent storage, return stored path."""
    settings = get_settings()
    dest_dir = settings.storage_dir / "documents"
    dest_dir.mkdir(parents=True, exist_ok=True)
    dest = dest_dir / f"{uuid4().hex}_{file_path.name}"
    shutil.copy2(file_path, dest)
    return dest


def ingest_document(file_path: Path, tender_id=None, bidder_id=None) -> tuple[Document, list[Page]]:
    """Ingest a single file, routing by type. Returns (Document, Pages)."""
    stored = _store_file(file_path)
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        document, pages = extract_pdf(stored, tender_id=tender_id, bidder_id=bidder_id)

        if document.routing_tag == RoutingTag.scanned_pdf:
            logger.info("Scanned PDF detected, running OCR: %s", file_path.name)
            pages = ocr_document_pages(pages, stored)

        return document, pages

    elif suffix in IMAGE_EXTENSIONS:
        img_bytes = stored.read_bytes()
        text, confidence = vision_page(img_bytes)

        from nirikshak.core.hashing import content_hash
        document = Document(
            id=uuid4(),
            tender_id=tender_id,
            bidder_id=bidder_id,
            filename=file_path.name,
            content_hash=content_hash(img_bytes),
            routing_tag=RoutingTag.photo_certificate,
        )
        page = Page(
            id=uuid4(),
            document_id=document.id,
            page_number=0,
            text=text,
            bboxes=[],
        )
        return document, [page]

    elif suffix == ".docx":
        from docx import Document as DocxDocument
        from nirikshak.core.hashing import content_hash

        raw_bytes = stored.read_bytes()
        docx = DocxDocument(stored)
        full_text = "\n".join(p.text for p in docx.paragraphs)

        document = Document(
            id=uuid4(),
            tender_id=tender_id,
            bidder_id=bidder_id,
            filename=file_path.name,
            content_hash=content_hash(raw_bytes),
            routing_tag=RoutingTag.native_pdf,
        )
        page = Page(
            id=uuid4(),
            document_id=document.id,
            page_number=0,
            text=full_text,
            bboxes=[],
        )
        return document, [page]

    else:
        raise ValueError(f"Unsupported file type: {suffix}")


def ingest_bidder_packet(bidder_id, tender_id, folder: Path) -> list[tuple[Document, list[Page]]]:
    """Ingest all files in a bidder's submission folder."""
    results = []
    for file_path in sorted(folder.iterdir()):
        if file_path.is_file() and not file_path.name.startswith("."):
            try:
                doc, pages = ingest_document(file_path, tender_id=tender_id, bidder_id=bidder_id)
                results.append((doc, pages))
                logger.info("Ingested bidder document: %s", file_path.name)
            except ValueError as e:
                logger.warning("Skipping unsupported file %s: %s", file_path.name, e)
    return results
